import dash
from dash.dependencies import Input, Output, State
import dash_html_components as html
import dash_core_components as dcc
import pandas as pd
import random
from datetime import date
import csv
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, send_from_directory
from urllib.parse import quote as urlquote
import os


UPLOAD_DIRECTORY = "documents"

server=Flask(__name__)
app = dash.Dash(server=server)


reader = csv.reader(open('Industries.csv', 'r'))
industries = {}
for row in reader:
    k, v = row
    industries[k] = v

reader2 = csv.reader(open('keywords.csv', 'r'))
keywords = {}
for row in reader2:
    j, k = row
    keywords[j] = k

def save_document(document, name):
    document.save('documents/{}_Profile.docx'.format(name))

@server.route("/download/<path:path>")
def download(path):
    """Serve a file from the upload directory."""
    return send_from_directory(UPLOAD_DIRECTORY, path, as_attachment=True)

app.layout = html.Div([

    html.H1('Big Data Profile Generator', style = { 'textAlign': 'center'}),

    html.Hr(),

    html.H2('Consultant Details',  style = { 'textAlign': 'center'}),
    html.Hr(),
    html.Table(
        [ html.Tr([html.Td(html.Label('Name')), html.Td(dcc.Input(id='input-1',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Phone Number')), html.Td(dcc.Input(id='input-2',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Email')), html.Td(dcc.Input(id='input-3',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Education - Name')), html.Td(dcc.Input(id='input-4',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Education - Location')), html.Td(dcc.Input(id='input-5',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Education - Degree')), html.Td(dcc.Input(id='input-6',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Certifications')), html.Td(dcc.Input(id='input-7',type='text', value='')) ]),
          html.Tr([html.Td(html.Label('Industry 1')), html.Td(dcc.Dropdown(
            id='dropdown-1', options=[{'label': industry, 'value': type} for industry, type in industries.items()], value=''))]),
          html.Tr([html.Td(html.Label('Industry 2')), html.Td(dcc.Dropdown(
            id='dropdown-2', options=[{'label': industry, 'value': type} for industry, type in industries.items()], value=''))]),
          html.Tr([html.Td(html.Label('Industry 3')), html.Td(dcc.Dropdown(
            id='dropdown-3', options=[{'label': industry, 'value': index} for industry, index in industries.items()], value=''))]),
          html.Tr([html.Td(html.Label('Tech Stack 1 (Select 8)')), html.Td(dcc.Dropdown(
            id='dropdown-4', options=[{'label': keyword, 'value': index} for keyword, index in keywords.items()], value='', multi=True))]),
          html.Tr([html.Td(html.Label('Tech Stack 2 (Select 8)')), html.Td(dcc.Dropdown(
            id='dropdown-5', options=[{'label': keyword, 'value': index} for keyword, index in keywords.items()], value='', multi=True))]),
          html.Tr([html.Td(html.Label('Tech Stack 3 (Select 8)')), html.Td(dcc.Dropdown(
            id='dropdown-6', options=[{'label': keyword, 'value': index} for keyword, index in keywords.items()], value='', multi=True))])
          ],


        style = { 'marginLeft': '40%', 'marginRight': '25%'} ),



    html.Hr(),
    html.Button('Generate', id='button-2',  style = { 'marginLeft': '48%', 'marginRight': '25%'}),
    html.Hr(),
    html.Button('Clear', id='button-3',  style = { 'marginLeft': '48%', 'marginRight': '25%'}),
    html.Hr(),
    html.Table([html.Tr([html.Td(html.Label('')), html.Td(html.Div(id='output'))])], style={ 'marginLeft': '45%', 'marginRight': '25%'}),
    html.Ul(id="file-list", style = { 'marginLeft': '48%', 'marginRight': '25%'}),
    html.H2('', id="file-list2"),
    html.H2('', id="file-list1"),
    dcc.Interval(id='interval-component', interval = 1 * 1000, n_intervals=5)

])

def file_download_link(filename):
    """Create a Plotly Dash 'A' element that downloads a file from the app."""
    location = "/download/{}".format(urlquote(filename))
    return html.A(filename, href=location)

def uploaded_files():
    """List the files in the upload directory."""
    files = []
    for filename in os.listdir(UPLOAD_DIRECTORY):
        path = os.path.join(UPLOAD_DIRECTORY, filename)
        if os.path.isfile(path):
            files.append(filename)
    return files

def delete_files():
    for file in os.scandir("documents"):
        if file.name.endswith(".docx"):
            os.unlink(file.path)
    return ''




@app.callback(Output('file-list', 'children'),
              [Input('interval-component', 'n_intervals')])
def update(n):
    files = uploaded_files()
    return [html.Li(file_download_link(filename)) for filename in files]

@app.callback(
    Output('file-list1', 'children'),
    [Input('button-2', 'n_clicks')],
    state=[State('input-1', 'value'),
     State('input-2', 'value'),
     State('input-3', 'value'),
     State('input-4', 'value'),
     State('input-5', 'value'),
     State('input-6', 'value'),
     State('input-7', 'value'),
     State('dropdown-1', 'value'),
     State('dropdown-2', 'value'),
     State('dropdown-3', 'value'),
     State('dropdown-4', 'value'),
     State('dropdown-5', 'value'),
     State('dropdown-6', 'value')
           ])


def compute(n_clicks, input1, input2, input3, input4, input5, input6, input7, dropdown1, dropdown2, dropdown3, dropdown4, dropdown5,dropdown6):
    df = pd.read_csv('Big_Data_Final.csv', index_col=0)
    df2 = pd.read_csv('Industry.csv', index_col = 0)
    # clients = random.sample(df['Client'].tolist(), k=3)
    clients0 = random.choice(df2[df2['keyword'].str.contains(dropdown1)]['Company'].tolist())
    clients1 = random.choice(df2[df2['keyword'].str.contains(dropdown2)]['Company'].tolist())
    clients2 = random.choice(df2[df2['keyword'].str.contains(dropdown3)]['Company'].tolist())
    loc1 = df.loc[df['Client'] == clients0].iloc[0]['Location']
    loc2 = df.loc[df['Client'] == clients1].iloc[0]['Location']
    loc3 = df.loc[df['Client'] == clients2].iloc[0]['Location']

    year = date.today().strftime("%Y")
    month = date.today().strftime("%B")

    bullets1=[]
    bullets2=[]
    bullets3=[]

    for key in dropdown4:
        bullets1+=random.sample(df[df['Job Description'].str.contains(key, case=False)]['Job Description'].tolist(), k=2)
    for key in dropdown5:
        bullets2+=random.sample(df[df['Job Description'].str.contains(key, case=False)]['Job Description'].tolist(), k=2)
    for key in dropdown6:
        bullets3 += random.sample(df[df['Job Description'].str.contains(key, case=False)]['Job Description'].tolist(), k=2)

    document = Document()

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = heading.add_run(input1)
    name.font.name = 'Calibri'
    name.font.size = Pt(30)

    info = heading.add_run('\n' + input2 + ' ' + input3)
    info.font.size = Pt(12)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section1 = document.add_heading(level=0)
    experience = section1.add_run('EXPERIENCE')
    experience.font.size = Pt(12)

    job1 = document.add_table(rows=1, cols=2)
    row1 = job1.rows[0].cells
    client1 = row1[0].paragraphs[0]
    client1.text = ''
    client1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    client1name = client1.add_run(clients0 + '\n')
    client1name.bold = True
    client1name.space_before = Pt(0)

    title1 = client1.add_run('Data Engineer')
    title1.italic = True

    location1 = row1[1].paragraphs[0]
    location1.text = ''
    location1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    location1name = location1.add_run(str(loc1) + '\n')
    location1name.space_before = Pt(0)

    dates1 = location1.add_run(month + ' ' + str(int(year) - 2) + ' - Present')

    for bullet in bullets1:
        document.add_paragraph(bullet, style='ListBullet')

    document.add_page_break()

    job2 = document.add_table(rows=1, cols=2)
    row2 = job2.rows[0].cells
    client2 = row2[0].paragraphs[0]
    client2.text = ''
    client2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    client2name = client2.add_run(clients1 + '\n')
    client2name.bold = True
    client2name.space_before = Pt(0)

    title2 = client2.add_run('Data Engineer')
    title2.italic = True

    location2 = row2[1].paragraphs[0]
    location2.text = ''
    location2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    location2name = location2.add_run(str(loc2) + '\n')
    location2name.space_before = Pt(0)

    dates2 = location2.add_run(month + ' ' + str(int(year) - 4) +' - ' + month + ' ' + str(int(year) - 2))

    for bullet in bullets2:
        document.add_paragraph(bullet, style='ListBullet')

    document.add_page_break()

    job3 = document.add_table(rows=1, cols=2)
    row3 = job3.rows[0].cells
    client3 = row3[0].paragraphs[0]
    client3.text = ''
    client3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    client3name = client3.add_run(clients2 + '\n')
    client3name.bold = True
    client3name.space_before = Pt(0)

    title3 = client3.add_run('Data Engineer')
    title3.italic = True

    location3 = row3[1].paragraphs[0]
    location3.text = ''
    location3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    location3name = location3.add_run(str(loc3) + '\n')
    location3name.space_before = Pt(0)

    dates3 = location3.add_run(month+ ' ' + str(int(year) - 6) +' - ' + month + ' ' + str(int(year) - 4))

    for bullet in bullets3:
        document.add_paragraph(bullet, style='ListBullet')

    document.add_page_break()

    section2 = document.add_heading(level=0)
    education = section2.add_run('EDUCATION')
    education.font.size = Pt(12)

    education_table = document.add_table(rows=1, cols=2)
    row1 = education_table.rows[0].cells
    school = row1[0].paragraphs[0]
    school.text = ''
    school.alignment = WD_ALIGN_PARAGRAPH.LEFT

    schoolname = school.add_run(str(input4) + '\n')
    schoolname.bold = True
    schoolname.space_before = Pt(0)

    degree = school.add_run(str(input6))
    degree.italic = True

    schlocation = row1[1].paragraphs[0]
    schlocation.text = ''
    schlocation.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    schlocationname = schlocation.add_run(str(input5) + '\n')
    schlocationname.space_before = Pt(0)

    return save_document(document, input1)

@app.callback(
    Output('file-list2', 'children'),
    [Input('button-3', 'n_clicks')])

def clear(n_clicks):
    return delete_files()

if __name__ == '__main__':
    app.run_server(port=8080)